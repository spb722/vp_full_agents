from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "vp_agent_processing_guide"
OUT_PATH = OUT_DIR / "Virtual_Profile_Agent_Processing_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF4EA"
PALE_AMBER = "FFF4CE"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "1F2937"


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    widths = [int(w) for w in widths]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in kwargs[edge]:
                element.set(qn(f"w:{key}"), str(kwargs[edge][key]))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def ensure_style(doc, name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = "Calibri Light"
    title.font.size = Pt(32)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(MID_GRAY)
    subtitle.paragraph_format.space_after = Pt(22)

    code = ensure_style(doc, "Code Block")
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(BLACK)
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.16)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.05

    callout = ensure_style(doc, "Callout")
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(4)
    callout.paragraph_format.line_spacing = 1.2

    small = ensure_style(doc, "Small Text")
    small.font.name = "Calibri"
    small.font.size = Pt(9)
    small.font.color.rgb = RGBColor.from_string(MID_GRAY)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def shade_paragraph(paragraph, fill=LIGHT_GRAY, border_color="D6DCE4"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "4")
        element.set(qn("w:color"), border_color)
        borders.append(element)


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p.add_run(text)
    shade_paragraph(p, LIGHT_GRAY)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "color": BLUE},
        start={"val": "single", "sz": 16, "color": BLUE},
        bottom={"val": "single", "sz": 8, "color": BLUE},
        end={"val": "single", "sz": 8, "color": BLUE},
    )
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout"]
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.add_run(text)
    return p


def new_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    existing_nums = numbering.findall(qn("w:num"))
    new_num_id = max([int(n.get(qn("w:numId"))) for n in existing_nums] or [0]) + 1

    # Reuse the standard decimal-list definition, but create a fresh instance
    # with an explicit level-0 restart at 1.
    abstract_num_id = "7"
    for num in existing_nums:
        if num.get(qn("w:numId")) == "5":
            ref = num.find(qn("w:abstractNumId"))
            if ref is not None:
                abstract_num_id = ref.get(qn("w:val"))
            break

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def add_number(doc, text, num_id):
    p = doc.add_paragraph(style="List Number")
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths, header_fill=PALE_BLUE, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
        run.font.size = Pt(font_size)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
    set_table_geometry(table, widths)
    for row in table.rows:
        row.height = None
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def keep_table_rows_together(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True


def add_step_flow(doc):
    rows = [
        ("1", "Receive", "Wrapper supplies client + marketer sentence; a request ID is created."),
        ("2", "Interpret", "Core skills guide semantic extraction; normalization produces evidence."),
        ("3", "Retrieve", "One batched call returns compact metric, filter, and time candidates."),
        ("4", "Resolve", "The agent chooses columns, table/date path, and the best seed or override."),
        ("5", "Render", "Only render_condition is allowed to emit PARENT_CONDITION syntax."),
        ("6", "Validate", "Deterministic checks confirm grammar, columns, table, and coverage."),
        ("7", "Return + audit", "The API returns the result; Langfuse retains the complete trace."),
    ]
    table = add_table(doc, ["", "Stage", "What happens"], rows, [600, 1560, 7200], font_size=9.2)
    for row in table.rows[1:]:
        set_cell_shading(row.cells[0], PALE_BLUE)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[0].paragraphs[0].runs[0].bold = True
    keep_table_rows_together(table)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.text = "VIRTUAL PROFILE AGENT  |  PROCESSING GUIDE"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)
    add_page_number(section.footer.paragraphs[0])

    # Editorial cover.
    doc.add_paragraph().paragraph_format.space_after = Pt(28)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    r = kicker.add_run("SYSTEM REFERENCE")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph("Virtual Profile Agent\nProcessing Guide", style="Title")
    doc.add_paragraph(
        "How a marketer’s sentence becomes an auditable, rendered, and validated PARENT_CONDITION",
        style="Subtitle",
    )

    add_callout(
        doc,
        "Core guarantee",
        "The language model interprets business meaning, but it never writes condition syntax by hand. "
        "A rule exists only after render_condition emits it and validate_rule accepts it.",
        PALE_GREEN,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    meta = add_table(
        doc,
        ["Applies to", "Workflow", "Worked example"],
        [["Omantel and Airtel VP requests", "Optimized orchestrator-led agent", "Real Omantel request 32e2594fe7"]],
        [3000, 3000, 3360],
        font_size=9.5,
    )
    keep_table_rows_together(meta)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph(style="Small Text")
    p.add_run("Document date: 15 July 2026  •  Audience: product, engineering, QA, and operations")
    doc.add_page_break()

    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_paragraph(
        "This guide explains the optimized Virtual Profile (VP) agent from request receipt to final API response. "
        "It describes what is loaded, what each skill and tool contributes, who owns each decision, what is kept "
        "in model context, and what is preserved in Langfuse for audit."
    )
    doc.add_paragraph(
        "The system is intentionally hybrid: the orchestrating model owns semantic interpretation and retry strategy, "
        "while deterministic tools provide evidence, enforce invariants, render syntax, and validate the result."
    )

    doc.add_heading("2. End-to-end workflow", level=1)
    add_step_flow(doc)
    add_callout(
        doc,
        "Important distinction",
        "Retrieval candidates and seed suggestions are traceable recommendations, not final truth. The agent may reject "
        "them when they conflict with the marketer’s meaning, but the override must be recorded.",
    )

    doc.add_heading("2.1 Request contract", level=2)
    doc.add_paragraph("The wrapper service sends two business inputs:")
    add_bullet(doc, "client — required and never inferred from the sentence.")
    add_bullet(doc, "sentence — the marketer’s audience or KPI request in plain language.")
    add_code(
        doc,
        '{\n  "client": "omantel",\n  "sentence": "Total pay-as-you-go data used on the local network by smartphone users over last 3 months"\n}',
    )
    doc.add_paragraph(
        "The service creates a request ID before orchestration. That ID links the API response, request-state audit, "
        "and Langfuse trace."
    )

    doc.add_heading("2.2 Stable ambiguity precheck", level=2)
    doc.add_paragraph(
        "Before expensive reasoning, stable business ambiguities may be detected. For example, “20% of recharge amount” "
        "could mean customers whose recharge is at least 20% of a benchmark, or a computed 20% share. When the meaning "
        "cannot be safely determined, the agent returns one plain-English clarification question instead of guessing."
    )

    doc.add_heading("3. Who owns each decision?", level=1)
    ownership = add_table(
        doc,
        ["Responsibility", "Owner", "Boundary"],
        [
            ("Business interpretation", "Orchestrating model", "Understands KPI, filters, aggregate, period, ambiguity, and comparison intent."),
            ("Candidate evidence", "Retrieval tools", "Surface ranked metadata; they do not declare final business truth."),
            ("Table/date routing", "Model + deterministic configuration", "The model chooses the semantic path; configuration supplies stable group date defaults."),
            ("Template proposal", "Seed selector", "Returns a traceable proposal and alternatives; the agent may promote a better alternative."),
            ("Condition syntax", "render_condition", "Exclusive syntax emitter. The model cannot patch or hand-write grammar."),
            ("Rule acceptance", "validate_rule", "Deterministically checks the emitted rule and reports errors or warnings."),
            ("Retry strategy", "Orchestrating model", "Retries retrieval/routing with corrected inputs; never repairs rendered syntax manually."),
        ],
        [2640, 2160, 4560],
        font_size=8.8,
    )
    keep_table_rows_together(ownership)

    doc.add_heading("4. Skill loading and progressive disclosure", level=1)
    doc.add_paragraph(
        "Skills are procedural instructions for the agent. The platform Skill capability loads a skill body only when "
        "its guidance is needed. Larger reference material remains outside the default prompt and is opened through Read "
        "only for the relevant case. This is progressive disclosure: keep the normal path compact, expand when uncertainty requires it."
    )

    doc.add_heading("4.1 Core skills", level=2)
    core_skills = add_table(
        doc,
        ["Skill", "When loaded", "What it teaches the agent"],
        [
            ("vp-extraction", "Every request", "How to decompose the sentence into KPI, domain, aggregate, filters, operator/value, and time intent."),
            ("vp-rendering-rules", "Every request", "The rendering boundary, placeholder handling, filter placement, grammar invariants, and validation requirements."),
        ],
        [2100, 1680, 5580],
        font_size=9,
    )
    keep_table_rows_together(core_skills)

    doc.add_heading("4.2 Conditional skills", level=2)
    conditional_skills = add_table(
        doc,
        ["Skill", "Trigger", "Purpose"],
        [
            ("vp-disambiguation", "Meaning remains materially ambiguous", "Builds one batched, plain-English clarification without exposing internal schema names."),
            ("vp-variant-selection", "A template/path must be chosen", "Selects Variant 1, 2, or 3 based on filters, tables, and period-comparison structure."),
            ("vp-metrics-comparison", "Two periods or metrics are compared", "Guides paired operands, aligned periods, comparison operators, and arithmetic structure."),
            ("vp-table-routing", "Candidates span groups or need a custom period", "Guides use of Customer 360 versus summary/event sources and the correct date-bearing group."),
            ("vp-golden-examples", "A reviewed semantic precedent helps", "Provides compact, approved examples for resolver/verifier alignment; it is not a phrase switchboard."),
        ],
        [2050, 2500, 4810],
        font_size=8.8,
    )
    keep_table_rows_together(conditional_skills)

    doc.add_heading("4.3 Reference material loaded only when needed", level=2)
    add_bullet(doc, "Time-token cases and calendar semantics, such as bounded last-N-month ranges versus a specific month offset.")
    add_bullet(doc, "Predicate/operator catalog for consistent business comparisons.")
    add_bullet(doc, "Golden-case patterns and reviewed semantic precedents.")
    add_bullet(doc, "Group/date defaults and exceptional event-date overrides.")
    add_bullet(doc, "Variant 3 examples for multi-period arithmetic and comparisons.")
    add_callout(
        doc,
        "Configured group date defaults",
        "Instant_cdr_group → FCT_DT; Common_Seg_Fct (Summary CDR) → COMMON_Event_date; "
        "Subscriptions → SUBSCRIPTIONS_DT; Recharge_Seg_Fct → RECHARGE_Event_Date; "
        "LIFECYCLE_CDR → L_SENT_DATE. SUBSCRIPTIONS_EVENT_DATE is reserved for renewal/cancellation semantics.",
        PALE_AMBER,
    )

    doc.add_heading("5. Tool loading and tool-call contract", level=1)
    doc.add_paragraph(
        "VP tool definitions are deferred to save prompt tokens. ToolSearch loads the required definitions by exact name. "
        "Once loaded, each tool call follows a consistent contract: compact input, compact operational output, agent interpretation, "
        "and a separate full audit record."
    )
    tool_contract_num = new_numbering_instance(doc)
    add_number(doc, "ToolSearch loads only the needed VP tool definitions.", tool_contract_num)
    add_number(doc, "The agent sends structured JSON input to the chosen tool.", tool_contract_num)
    add_number(doc, "The tool returns deterministic evidence or an enforced result.", tool_contract_num)
    add_number(doc, "Compact fields enter model context; complete rankings and internals go to Langfuse/request state.", tool_contract_num)
    add_number(doc, "The agent continues, retries, clarifies, or stops according to the tool’s authority.", tool_contract_num)

    doc.add_heading("5.1 Active VP tools", level=2)
    tools_table = add_table(
        doc,
        ["Tool", "Typical input", "Output used by agent", "Authority"],
        [
            ("normalize_slots", "client + sentence", "Initial slots, ambiguity flags, normalization warnings", "Evidence; agent may correct semantic misses"),
            ("retrieve_columns", "client + corrected slots in one batch", "Up to five compact candidates per role, scores, evidence, audit ID", "Evidence; full ranking remains out of context"),
            ("select_seed", "client + audit ID/slots; optional seed ID", "One complete selected seed + three compact alternatives", "Proposal; deterministic promotion is allowed"),
            ("retrieve_existing_vps", "comparison intent/columns", "Relevant historical VP conditions", "Evidence for complex Variant 3 structure"),
            ("record_resolution", "chosen columns, seed, path, slots, override", "Audit confirmation", "Records the decision; does not make it"),
            ("render_condition", "complete template + variables + filters", "Final PARENT_CONDITION syntax", "Exclusive syntax emitter"),
            ("validate_rule", "rule + request + table", "ok/errors/warnings/referenced columns", "Deterministic acceptance gate"),
        ],
        [2100, 2300, 3000, 1960],
        font_size=8.2,
    )
    keep_table_rows_together(tools_table)

    doc.add_heading("5.2 Supporting platform capabilities", level=2)
    platform = add_table(
        doc,
        ["Capability", "Function"],
        [
            ("Skill", "Loads procedural guidance for the current reasoning stage."),
            ("Read", "Opens a targeted reference only when the skill directs the agent to it."),
            ("ToolSearch", "Loads deferred tool definitions instead of placing every schema in the initial context."),
            ("Agent/verifier", "Optional for low-confidence cases; required for Variant 3 period comparisons before final acceptance."),
        ],
        [1900, 7460],
        font_size=9,
    )
    keep_table_rows_together(platform)

    doc.add_heading("6. Context versus audit data", level=1)
    split = add_table(
        doc,
        ["Model context: compact operational view", "Langfuse/request state: complete audit view"],
        [
            ("Top candidates by role: feature name, group, type, short description, time-window support, score, concise evidence.",
             "Full BM25, embedding, semantic, metadata, prior, and hybrid rankings with all candidates."),
            ("One complete selected seed and three compact alternatives.",
             "Full seed ranking, signatures, confidence components, proposal/override history."),
            ("Current resolution state and validation result.",
             "Every skill/tool call, input/output, latency, token usage, model cost, request ID, and final response."),
        ],
        [4680, 4680],
        font_size=8.8,
    )
    keep_table_rows_together(split)
    add_callout(
        doc,
        "Why this matters",
        "Token reduction comes from removing diagnostic detail from the reasoning context—not from removing evidence or auditability. "
        "If confidence is low or group coverage is missing, the agent deterministically expands retrieval.",
        PALE_GREEN,
    )

    doc.add_page_break()
    doc.add_heading("7. Worked example — real Omantel request", level=1)
    doc.add_paragraph(
        "This example is taken from the optimized 34-case regression run and its corrected Langfuse export. It shows a useful "
        "property of the architecture: the first deterministic suggestion can be imperfect, while the final result remains correct "
        "because evidence, semantic reasoning, rendering, and validation have separate responsibilities."
    )

    example_overview = add_table(
        doc,
        ["Field", "Value"],
        [
            ("Client", "omantel"),
            ("Request ID", "32e2594fe7"),
            ("Sentence", "Total pay-as-you-go data used on the local network by smartphone users over last 3 months"),
            ("Expected meaning", "SUM of local PAYG data volume, smartphone filter, bounded previous three complete months"),
        ],
        [1800, 7560],
        font_size=9.2,
    )
    keep_table_rows_together(example_overview)

    doc.add_heading("7.1 Core guidance is loaded", level=2)
    doc.add_paragraph(
        "The trace shows vp-extraction and vp-rendering-rules being loaded first. Variant-selection and golden-example guidance "
        "is then loaded for path and semantic alignment. The platform also loads the required VP tool definitions through ToolSearch."
    )

    doc.add_heading("7.2 Initial normalization produces evidence", level=2)
    add_code(
        doc,
        '{\n  "client": "omantel",\n  "request": "Total pay-as-you-go data used on the local network by smartphone users over last 3 months"\n}',
    )
    doc.add_paragraph(
        "normalize_slots recognized the smartphone filter and a three-month period, but initially returned an unknown domain, "
        "an empty KPI phrase, and a clarification warning. That output is not allowed to become the final interpretation automatically."
    )
    add_callout(
        doc,
        "Agent correction",
        "From the sentence, the agent resolves domain = usage, KPI = local pay-as-you-go data volume, aggregate = SUM, "
        "filter = handset type Smartphone, and time = bounded last three complete months.",
    )
    add_code(
        doc,
        '{\n  "domain": "usage",\n  "kpi_phrase": "pay-as-you-go data usage - local network",\n  "aggregate": "SUM",\n  "filters": [{"phrase": "handset type - smartphone", "operator": "=", "value": "Smartphone"}],\n  "time_token": "M3 (bounded last-3-months range)"\n}',
    )

    doc.add_heading("7.3 One batched retrieval resolves all roles", level=2)
    doc.add_paragraph(
        "The corrected slots are passed once to retrieve_columns. The call resolves the metric, filter, and time requirements together, "
        "which reduces repeated tool prompts and makes cross-role/group compatibility visible."
    )
    retrieval = add_table(
        doc,
        ["Role", "Best compact candidate", "Group", "Score", "Concise evidence"],
        [
            ("Metric", "COMMON_Data_Local_PayG_Volume", "Common_Seg_Fct", "0.937", "Exact local PAYG data-volume meaning; supports custom rolling windows via the group date."),
            ("Filter", "Profile_Cdr_Handset_Type", "Profile_Cdr_group", "1.065", "Direct categorical handset-type field; matches Smartphone."),
            ("Time", "COMMON_Event_Date", "Common_Seg_Fct", "—", "Configured Summary CDR event date for bounded usage windows."),
        ],
        [860, 3050, 1700, 850, 2900],
        font_size=8.1,
    )
    keep_table_rows_together(retrieval)
    doc.add_paragraph(
        "The response includes audit ID 0293428870c9. The model receives only compact candidates; the complete ranking, including "
        "BM25/embedding components and lower-ranked candidates, stays in Langfuse."
    )

    doc.add_heading("7.4 Table routing is resolved", level=2)
    doc.add_paragraph(
        "The metric and filter originate from different groups. vp-table-routing guides a multi-table filter-then-aggregate path. "
        "The aggregate is evaluated on Common_Seg_Fct (Summary CDR), whose configured date field is COMMON_Event_Date."
    )

    doc.add_heading("7.5 The first seed proposal is checked—not blindly accepted", level=2)
    seed_table = add_table(
        doc,
        ["Seed", "Meaning", "Confidence", "Decision"],
        [
            ("S01_fixed_month_sum", "SUM for one specific calendar-month offset", "0.759", "Rejected: M3 alone would mean only the month three months ago."),
            ("S13_last_n_months_bounded", "SUM across a bounded N-month range", "0.752", "Promoted: matches the previous three complete months."),
        ],
        [2100, 3460, 1100, 2700],
        font_size=8.8,
    )
    keep_table_rows_together(seed_table)
    add_callout(
        doc,
        "Why the lower score wins",
        "A 0.007 score difference cannot override clear time semantics. “Last 3 months” is a range: date >= CurrentMonth-3MONTHS "
        "and date < CurrentMonth. The override is recorded for audit.",
        PALE_AMBER,
    )

    doc.add_heading("7.6 Resolution is recorded", level=2)
    doc.add_paragraph("record_resolution stores the chosen interpretation before rendering:")
    resolution_rows = [
        ("Columns", "COMMON_Data_Local_PayG_Volume; Profile_Cdr_Handset_Type; COMMON_Event_Date"),
        ("Seed", "S13_last_n_months_bounded"),
        ("Proposed seed", "S01_fixed_month_sum"),
        ("Override", "true"),
        ("Path", "variant2_multi_table_filter_then_aggregate"),
        ("Snapshot", "false"),
    ]
    resolution = add_table(doc, ["Recorded item", "Value"], resolution_rows, [1900, 7460], font_size=8.8)
    keep_table_rows_together(resolution)

    doc.add_heading("7.7 The renderer emits the rule", level=2)
    doc.add_paragraph(
        "The agent supplies a complete template to render_condition. The renderer—not the model—emits the authoritative syntax. "
        "Runtime placeholders remain intact for the downstream UI."
    )
    add_code(
        doc,
        "Profile_Cdr_Handset_Type = 'Smartphone' AND COMMON_Event_Date >= CurrentMonth-3MONTHS AND "
        "COMMON_Event_Date < CurrentMonth AND SUM(COMMON_Data_Local_PayG_Volume) ${operator} ${value}",
    )

    doc.add_heading("7.8 Deterministic validation accepts the result", level=2)
    add_code(
        doc,
        '{\n  "ok": true,\n  "table": "Common_Seg_Fct",\n  "errors": [],\n  "warnings": [],\n  "referenced_columns": [\n    "COMMON_Data_Local_PayG_Volume",\n    "COMMON_Event_Date",\n    "Profile_Cdr_Handset_Type"\n  ],\n  "coverage_terms": ["data", "local", "smartphone"]\n}',
    )
    doc.add_paragraph(
        "Because validation is successful, no retry or clarification is needed. This low-risk Variant 2 case does not require a verifier subagent; "
        "a verifier becomes important for low-confidence resolutions and is mandatory for Variant 3 period comparisons."
    )

    doc.add_heading("7.9 Final API response", level=2)
    add_code(
        doc,
        '{\n  "ok": true,\n  "mode": "agent",\n  "request_id": "32e2594fe7",\n  "parent_condition": "Profile_Cdr_Handset_Type = \'Smartphone\' AND COMMON_Event_Date >= CurrentMonth-3MONTHS AND COMMON_Event_Date < CurrentMonth AND SUM(COMMON_Data_Local_PayG_Volume) ${operator} ${value}",\n  "selected_columns": [\n    "COMMON_Data_Local_PayG_Volume",\n    "Profile_Cdr_Handset_Type",\n    "COMMON_Event_Date"\n  ],\n  "seed_id": "S13_last_n_months_bounded",\n  "path": "variant2_multi_table_filter_then_aggregate",\n  "validation": {"ok": true, "errors": [], "warnings": []}\n}',
    )

    doc.add_heading("8. How the final result is reached", level=1)
    doc.add_paragraph(
        "The final rule is not the output of one prompt or one similarity score. It is the outcome of a staged decision chain with explicit checks:"
    )
    result_num = new_numbering_instance(doc)
    add_number(doc, "The sentence is decomposed into business roles and possible ambiguities.", result_num)
    add_number(doc, "Deterministic normalization supplies initial evidence; semantic gaps are corrected by the agent.", result_num)
    add_number(doc, "Batched retrieval supplies compact metric, filter, and time candidates with scores and evidence.", result_num)
    add_number(doc, "Routing selects the appropriate data group and configured date column.", result_num)
    add_number(doc, "Seed selection proposes a reusable structure; the agent checks it against the exact time and comparison meaning.", result_num)
    add_number(doc, "The decision is recorded, including any override and its traceable inputs.", result_num)
    add_number(doc, "render_condition alone emits the condition syntax.", result_num)
    add_number(doc, "validate_rule accepts the rule or returns a deterministic failure class.", result_num)
    add_number(doc, "The API returns the compact result while Langfuse retains the complete execution record.", result_num)

    doc.add_heading("9. Failure, clarification, and retry behavior", level=1)
    retries = add_table(
        doc,
        ["Situation", "System behavior"],
        [
            ("Business ambiguity", "Ask one batched, plain-English clarification question; do not expose internal schema names."),
            ("Column or coverage failure", "Retrieve again with the failed column excluded and inspect the next compact candidates."),
            ("Low confidence or missing group coverage", "Deterministically expand beyond the default five candidates for the affected role."),
            ("Routing or date failure", "Route again using corrected semantic slots and the configured group/date map."),
            ("Seed mismatch", "Promote a more appropriate alternative and record proposed-versus-selected seed."),
            ("Render or grammar failure", "Stop as a code defect. The model is never asked to patch syntax manually."),
            ("Variant 3 comparison", "Use comparison guidance and require verifier review before final acceptance."),
        ],
        [2700, 6660],
        font_size=8.8,
    )
    keep_table_rows_together(retries)

    doc.add_heading("10. Operational guarantees", level=1)
    add_bullet(doc, "Client is mandatory and never inferred from the marketer’s sentence.")
    add_bullet(doc, "The model owns meaning; deterministic components own evidence, invariants, rendering, and validation.")
    add_bullet(doc, "Only render_condition emits PARENT_CONDITION syntax.")
    add_bullet(doc, "Retrieval and seed selection are auditable recommendations, not irreversible decisions.")
    add_bullet(doc, "Compact context reduces tokens without deleting the complete ranking or decision history.")
    add_bullet(doc, "Every completed request can be joined across API logs, request state, and Langfuse using request_id.")
    add_bullet(doc, "Clarifications are written for marketers; internal table, group, and column names stay hidden unless explicitly requested.")

    add_callout(
        doc,
        "Summary",
        "The optimized VP agent is best understood as an auditable decision pipeline: skills guide interpretation, tools expose compact evidence, "
        "the orchestrator resolves meaning, the renderer emits syntax, and the validator determines whether the rule is safe to return.",
        PALE_GREEN,
    )

    # Keep table rows from splitting and add document metadata.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    props = doc.core_properties
    props.title = "Virtual Profile Agent Processing Guide"
    props.subject = "VP agent skills, tools, rendering, validation, and worked Omantel example"
    props.author = "VP Agent Team"
    props.keywords = "virtual profile, Omantel, agent, skills, tools, Langfuse, audit"

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
